import os
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Directorio
dir_path = r"G:\Unidades compartidas\GM_OP_ANALISIS_DE_DATOS\21.1 validacion kms\RESPUESTA_UF17\20230604"

# Encuentra el archivo .csv en el directorio que contiene 'resumen.csv' en su nombre
for file in os.listdir(dir_path):
    if 'resumen.csv' in file:
        csv_file_path = os.path.join(dir_path, file)
        break

# Leer el archivo .csv y seleccionar columnas específicas
df = pd.read_csv(csv_file_path)
selected_columns = ['fecha_viaje', 'id_fase_v', 'km_a_reclamar', 'respuesta', 'servicio_retoma', 'servicio', 'id_viaje', 'linea_sae', 'vehiculo']
lista_objeciones = df[selected_columns]

# Crear documentos Word a partir de la columna 'respuesta'
for i in range(len(lista_objeciones)):
    doc = Document()
    title = doc.add_heading('Resumen', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = doc.add_paragraph('Fecha Viaje: ' + str(lista_objeciones.loc[i, 'fecha_viaje']))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = doc.add_paragraph('ID Fase V: ' + str(lista_objeciones.loc[i, 'id_fase_v']))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = doc.add_paragraph('Servicio: ' + str(lista_objeciones.loc[i, 'servicio']))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = doc.add_paragraph('ID Viaje: ' + str(lista_objeciones.loc[i, 'id_viaje']))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = doc.add_paragraph('Linea SAE: ' + str(lista_objeciones.loc[i, 'linea_sae']))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph = doc.add_paragraph('Vehiculo: ' + str(lista_objeciones.loc[i, 'vehiculo']))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    observation = doc.add_paragraph('Observación: ' + lista_objeciones.loc[i, 'respuesta'])
    observation.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc_file_path = os.path.join(dir_path, f"{lista_objeciones.loc[i, 'id_fase_v']}.docx")
    doc.save(doc_file_path)

    # Agregar la ruta del archivo doc a la columna 'ruta' en 'lista_objeciones'
    lista_objeciones.loc[i, 'ruta'] = doc_file_path

# Ordenar el DataFrame por la columna 'km_a_reclamar' de mayor a menor
lista_objeciones = lista_objeciones.sort_values(by='km_a_reclamar', ascending=False)

# Escribir el DataFrame a un archivo .csv
csv_output_file_path = os.path.join(dir_path, 'resumen_respuestas.csv')
lista_objeciones.to_csv(csv_output_file_path, index=False)

print(lista_objeciones)
print("Documento Word Generado")